import io
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ─────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1E, 0x3A, 0x5F)
MID_BLUE   = RGBColor(0x2E, 0x86, 0xC1)
LIGHT_BLUE = RGBColor(0xD6, 0xEA, 0xF8)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT  = RGBColor(0x2C, 0x3E, 0x50)
GRAY_BG    = RGBColor(0xF2, 0xF3, 0xF4)


def _hex(color: RGBColor) -> str:
    return f"{color.rgb:06X}"


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, hex_color="CCCCCC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _para(doc, text, bold=False, size=11, color=DARK_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def _heading(doc, text, level=1):
    colors = {1: DARK_BLUE, 2: MID_BLUE, 3: DARK_TEXT}
    sizes  = {1: 18, 2: 14, 3: 12}
    p = _para(doc, text, bold=True, size=sizes[level], color=colors[level],
              space_before=12, space_after=6)
    # bottom border for h1 / h2
    if level <= 2:
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), _hex(MID_BLUE))
        pBdr.append(bot)
        pPr.append(pBdr)
    return p


def _add_table(doc, df: pd.DataFrame, max_rows=50):
    """Add a fully styled table from a DataFrame."""
    display = df.head(max_rows).copy()

    # Format numbers
    for col in display.select_dtypes(include="number").columns:
        if pd.api.types.is_integer_dtype(display[col]):
            display[col] = display[col].apply(lambda x: f"{int(x):,}" if pd.notna(x) else "—")
        else:
            display[col] = display[col].apply(lambda x: f"{x:,.2f}" if pd.notna(x) else "—")
    for col in display.select_dtypes(include=["datetime64"]).columns:
        display[col] = display[col].apply(lambda x: str(x.date()) if pd.notna(x) else "—")
    display = display.fillna("—").astype(str)

    tbl = doc.add_table(rows=1 + len(display), cols=len(display.columns))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # Header row
    hdr = tbl.rows[0]
    for i, col in enumerate(display.columns):
        cell = hdr.cells[i]
        _set_cell_bg(cell, _hex(DARK_BLUE))
        _set_cell_border(cell, "FFFFFF")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(str(col).upper())
        run.bold = True
        run.font.size  = Pt(9)
        run.font.color.rgb = WHITE

    # Data rows
    for r_idx, (_, row) in enumerate(display.iterrows()):
        bg = _hex(GRAY_BG) if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            _set_cell_bg(cell, bg)
            _set_cell_border(cell, "DDDDDD")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_TEXT

    # Column widths — spread evenly across page
    col_count = len(display.columns)
    col_width  = Inches(6.5 / col_count)
    for row in tbl.rows:
        for cell in row.cells:
            cell.width = col_width

    doc.add_paragraph()   # spacing after table


def _cover_page(doc, title: str, subtitle: str):
    # Blue banner
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, _hex(DARK_BLUE))
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run("🔍  InsightIQ")
    run.bold           = True
    run.font.size      = Pt(26)
    run.font.color.rgb = WHITE

    p2 = cell.add_paragraph(title)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after  = Pt(4)
    r2 = p2.runs[0]
    r2.bold           = True
    r2.font.size      = Pt(16)
    r2.font.color.rgb = RGBColor(0xA9, 0xCC, 0xE3)

    p3 = cell.add_paragraph(subtitle)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after  = Pt(20)
    r3 = p3.runs[0]
    r3.font.size      = Pt(10)
    r3.font.color.rgb = RGBColor(0x7F, 0xB3, 0xD3)

    doc.add_paragraph()


def _summary_box(doc, stats: dict):
    """Render key metric boxes as a 1-row table."""
    if not stats:
        return
    cols  = list(stats.items())[:4]
    tbl   = doc.add_table(rows=2, cols=len(cols))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(cols):
        # Value cell
        vc = tbl.rows[0].cells[i]
        _set_cell_bg(vc, _hex(DARK_BLUE))
        _set_cell_border(vc, "FFFFFF")
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vp.paragraph_format.space_before = Pt(10)
        vp.paragraph_format.space_after  = Pt(4)
        vr = vp.add_run(str(value))
        vr.bold = True; vr.font.size = Pt(20); vr.font.color.rgb = WHITE
        # Label cell
        lc = tbl.rows[1].cells[i]
        _set_cell_bg(lc, _hex(MID_BLUE))
        _set_cell_border(lc, "FFFFFF")
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(10)
        lr = lp.add_run(label.upper())
        lr.font.size = Pt(8); lr.font.color.rgb = WHITE
    for row in tbl.rows:
        for cell in row.cells:
            cell.width = Inches(6.5 / len(cols))
    doc.add_paragraph()


def generate_report(
    title: str,
    sections: list,           # [{"heading": str, "explanation": str, "df": DataFrame, "sql": str}]
    summary_stats: dict = None,
    username: str = "User",
) -> bytes:
    """
    Build a .docx report and return its bytes.

    sections = list of dicts:
        heading     : section title
        explanation : plain-text summary from Claude
        df          : DataFrame (optional)
        sql         : SQL used (optional, shown collapsed)
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Cover
    generated_at = datetime.now().strftime("%B %d, %Y  %H:%M")
    _cover_page(doc, title, f"Generated by {username}  •  {generated_at}  •  DataWarehouseAnalytics")

    # Summary stats bar
    if summary_stats:
        _heading(doc, "📊 Key Metrics at a Glance", level=1)
        _summary_box(doc, summary_stats)

    # Sections
    for sec in sections:
        _heading(doc, sec.get("heading", "Section"), level=2)

        explanation = sec.get("explanation", "")
        if explanation:
            for line in explanation.split("\n"):
                line = line.strip()
                if not line:
                    continue
                # Render bullet lines
                if line.startswith(("- ", "• ", "* ")):
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after  = Pt(2)
                    run = p.add_run(line.lstrip("-•* "))
                    run.font.size = Pt(10)
                    run.font.color.rgb = DARK_TEXT
                elif line.startswith("**") and line.endswith("**"):
                    _para(doc, line.strip("*"), bold=True, size=10, space_before=4, space_after=2)
                else:
                    _para(doc, line, size=10, space_before=2, space_after=2)

        df = sec.get("df")
        if df is not None and not df.empty:
            _para(doc, f"  {len(df)} rows returned", size=9,
                  color=RGBColor(0x71, 0x80, 0x96), space_before=6, space_after=4)
            _add_table(doc, df)

        sql = sec.get("sql", "")
        if sql:
            _para(doc, "SQL Query Used:", bold=True, size=9,
                  color=RGBColor(0x71, 0x80, 0x96), space_before=8, space_after=2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(8)
            run = p.add_run(sql)
            run.font.name  = "Courier New"
            run.font.size  = Pt(8)
            run.font.color.rgb = RGBColor(0x63, 0xB3, 0xED)

    # Footer on every page
    for sec in doc.sections:
        footer_para = sec.footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.clear()
        run = footer_para.add_run(
            f"InsightIQ Report  •  {title}  •  {generated_at}  •  Confidential"
        )
        run.font.size      = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
